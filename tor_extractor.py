"""
tor_extractor.py — Task 4

Extracts structured requirements from a Terms of Reference (ToR) document.
Supports .docx and .pdf file types.
"""
import io
import json
import logging
import re

import anthropic

from config import MODEL_NAME, ANTHROPIC_API_KEY

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Custom exception (available for callers; never raised inside extract_tor)
# ---------------------------------------------------------------------------

class ExtractionError(Exception):
    """Raised by callers who want to surface an error-state dict as an exception."""
    pass


# ---------------------------------------------------------------------------
# Prompts (from SPEC.md §6.1)
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = (
    "You are analyzing a Terms of Reference or Notice of Funding Opportunity document. "
    "Extract structured information accurately.\n"
    "Use the document's own language — do not normalize, generalize, or substitute terms. "
    "If a term appears explicitly in the document, use it exactly as written.\n"
    "Return ONLY valid JSON. No markdown fences. No explanation."
)

USER_PROMPT_TEMPLATE = (
    "Analyze this document and return a JSON object with these fields:\n\n"
    '"title": The title or name of the opportunity. String.\n'
    '"funder": The organization or government body funding this. '
    "Use exact name as written. String.\n"
    '"geography": List of countries or regions that are the TARGET or PRIORITY of this '
    "opportunity. Use names exactly as written in the document. Do not include countries "
    "mentioned only as context or comparisons. List of strings.\n"
    '"thematic_areas": List of themes, technical areas, or subject areas that this '
    "opportunity is about. Use the document's own terms — do not translate or generalize "
    "(e.g. if the doc says \"Financial Intelligence\" keep that, do not replace with "
    "\"AML\"). List of strings.\n"
    '"key_requirements": List of what the responding organization must demonstrate, '
    "deliver, or possess. Each item is one requirement. List of strings. Maximum 8 items.\n"
    '"evaluation_criteria": List of how responses will be evaluated. List of strings. '
    "Maximum 5 items. Empty list if not stated.\n"
    '"language": "English" or "Spanish".\n'
    '"extraction_confidence": "HIGH" if you found all main fields clearly stated. '
    '"MEDIUM" if some fields required inference. '
    '"LOW" if significant information was missing or ambiguous.\n'
    '"paragraphs": List of all paragraphs from the document in order. Each paragraph is '
    "a string. Exclude empty paragraphs and paragraphs under 20 characters. "
    "Maximum 80 paragraphs.\n"
    '"source_map": For each item in geography, thematic_areas, funder, and '
    "key_requirements, return the paragraph index (0-based position in your paragraphs "
    "list) and a short snippet (max 25 words) from that paragraph that contains the item. "
    'Structure: {{"geography": [{{"term": ..., "paragraph_index": ..., "snippet": ...}}], '
    '"thematic_areas": [...], "funder": [...], "key_requirements": [...]}} '
    "If the paragraph cannot be identified, use paragraph_index: -1 and snippet: \"\".\n\n"
    "DOCUMENT TEXT:\n"
    "{tor_text}"
)

# Required keys and their default values
_REQUIRED_KEYS = {
    "title": "",
    "funder": "",
    "geography": [],
    "thematic_areas": [],
    "key_requirements": [],
    "evaluation_criteria": [],
    "language": "English",
    "source_file": "",
    "extraction_confidence": "LOW",
    "paragraphs": [],
    "source_map": {
        "geography": [],
        "thematic_areas": [],
        "funder": [],
        "key_requirements": [],
    },
}


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _error_state(filename: str, message: str) -> dict:
    """Return a well-formed error-state dict."""
    return {
        "title": "",
        "funder": "",
        "geography": [],
        "thematic_areas": [],
        "key_requirements": [],
        "evaluation_criteria": [],
        "language": "English",
        "source_file": filename,
        "extraction_confidence": "LOW",
        "error": message,
    }


def _extract_text_docx(file_bytes: bytes) -> str:
    """Extract all text from a .docx file, including table cells."""
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def _extract_text_pdf(file_bytes: bytes) -> str:
    """Extract all text from a .pdf file using pdfplumber."""
    import pdfplumber

    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)

    return "\n".join(parts)


def _truncate_text(text: str) -> str:
    """
    Truncate text to at most 15,000 characters.
    Preserves first 12,000 + last 3,000 characters.
    """
    if len(text) > 15000:
        return text[:12000] + text[-3000:]
    return text


def _call_claude(client: anthropic.Anthropic, user_prompt: str) -> str:
    """Make a single Claude API call and return the raw text response."""
    response = client.messages.create(
        model=MODEL_NAME,
        max_tokens=8000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_prompt}],
    )
    return response.content[0].text


def _strip_markdown_fences(text: str) -> str:
    text = text.strip()
    # Find the first { and last } and extract only the JSON object
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        return text[start:end + 1]
    return text


def _fill_defaults(data: dict, filename: str) -> dict:
    """
    Ensure all required keys are present in data.
    Missing keys are filled with defaults. source_file is always overridden.
    """
    for key, default in _REQUIRED_KEYS.items():
        if key not in data:
            # Use a fresh copy for mutable defaults
            data[key] = default.copy() if isinstance(default, (dict, list)) else default
    # Ensure source_map has all four sub-keys
    sm = data.get("source_map", {})
    if not isinstance(sm, dict):
        sm = {}
    for sub_key in ("geography", "thematic_areas", "funder", "key_requirements"):
        if sub_key not in sm:
            sm[sub_key] = []
    data["source_map"] = sm
    # Always override source_file with the actual filename
    data["source_file"] = filename
    return data


# ---------------------------------------------------------------------------
# Public interface
# ---------------------------------------------------------------------------

def extract_tor(file_bytes: bytes, filename: str) -> dict:
    """
    Extract structured requirements from a ToR document.
    Returns a tor_data dict matching the TorData schema.
    On error, returns an error-state dict with extraction_confidence="LOW".
    Never raises exceptions — all errors are returned as dicts.
    """
    # 4.2 / 4.7 — File type guard
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in ("docx", "pdf"):
        return _error_state(filename, "Unsupported file type. Please upload a .docx or .pdf file.")

    # 4.2 — Text extraction
    try:
        if ext == "docx":
            text = _extract_text_docx(file_bytes)
        else:
            text = _extract_text_pdf(file_bytes)
    except Exception as exc:
        return _error_state(filename, f"Failed to extract text from file: {exc}")

    # 4.3 — Text length guard
    if len(text) < 100:
        logger.warning(
            "Extracted text is very short (%d chars) for file '%s'. Proceeding anyway.",
            len(text),
            filename,
        )

    # 4.3 — Truncation
    text = _truncate_text(text)

    # 4.4 / 4.5 — Claude API call
    user_prompt = USER_PROMPT_TEMPLATE.format(tor_text=text)

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    except Exception as exc:
        return _error_state(filename, f"Failed to initialise Anthropic client: {exc}")

    # First attempt
    try:
        raw_text = _call_claude(client, user_prompt)
    except Exception as exc:
        return _error_state(filename, f"Claude API call failed: {exc}")

    # 4.5 — JSON parsing with one retry
    try:
        data = json.loads(_strip_markdown_fences(raw_text))
    except json.JSONDecodeError:
        # Retry once
        try:
            raw_text = _call_claude(client, user_prompt)
        except Exception as exc:
            return _error_state(filename, f"Claude API call failed on retry: {exc}")

        try:
            data = json.loads(_strip_markdown_fences(raw_text))
        except json.JSONDecodeError:
            return _error_state(
                filename,
                "Claude API returned invalid JSON after retry. Please try again.",
            )

    # 4.5 — Fill missing keys with defaults, override source_file
    data = _fill_defaults(data, filename)

    return data
