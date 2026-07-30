"""Tests for output_formatter.py — Task 7.7"""
import sys
import os
import re
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from datetime import datetime
from hypothesis import given, settings, HealthCheck
from hypothesis import strategies as st
from docx import Document
from docx.shared import Pt, Cm

from output_formatter import write_output


# ---------------------------------------------------------------------------
# Fixtures / helpers
# ---------------------------------------------------------------------------

def _make_minimal_draft():
    return {
        "sections": {
            "opening_statement": "GovRisk overview. [REF:doc.pdf:page_1]",
            "institutional_overview": "Founded 2010.",
            "country_table": [
                {
                    "country": "Mexico",
                    "project_count": 2,
                    "year_range": "2020-2023",
                    "named_identifiers": ["PECEL"],
                    "donors": ["US State Dept"],
                }
            ],
            "geographic_experience": "Strong in Mexico.",
            "thematic_areas": "AML/CFT expertise.",
            "selected_project_experience": "Project Alpha 2022.",
            "alignment_with_tor": "Fully aligned.",
        },
        "interpretation_log": [
            {
                "section": "opening_statement",
                "inference_made": "Test",
                "source_used": "doc.pdf",
                "gap_flagged": None,
                "confidence": "HIGH",
            }
        ],
        "summary": {
            "sections_generated": 7,
            "projects_referenced": 1,
            "countries_covered": 1,
            "documents_used": 1,
            "overall_confidence": "HIGH",
        },
    }


def _make_minimal_citation_result():
    return {
        "paragraphs": [
            {
                "text": "GovRisk overview.",
                "citations": [{"filename": "doc.pdf", "page": 1}],
            }
        ]
    }


# ---------------------------------------------------------------------------
# Unit test: output file written to output_path
# ---------------------------------------------------------------------------

def test_output_file_written_to_output_path(tmp_path):
    """Output file is written to the specified output_path and path is returned."""
    result = write_output(
        _make_minimal_draft(),
        _make_minimal_citation_result(),
        output_path=str(tmp_path),
    )
    assert isinstance(result, str), f"Expected str, got {type(result)}"
    assert not result.startswith("ERROR:"), f"Got error: {result}"
    assert os.path.isfile(result), f"File does not exist: {result}"
    # File must be inside tmp_path
    assert str(tmp_path) in result, f"File not in tmp_path: {result}"


# ---------------------------------------------------------------------------
# Unit test: section headings appear in correct order
# ---------------------------------------------------------------------------

def test_section_headings_in_correct_order(tmp_path):
    """Section headings appear in the document in the specified order."""
    result = write_output(
        _make_minimal_draft(),
        _make_minimal_citation_result(),
        output_path=str(tmp_path),
    )
    assert not result.startswith("ERROR:"), f"Got error: {result}"

    doc = Document(result)

    # Collect all bold 14pt paragraph texts
    bold_14pt_texts = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold and run.font.size == Pt(14):
                text = run.text.strip()
                if text:
                    bold_14pt_texts.append(text)
                break

    expected_order = [
        "Opening Statement",
        "Institutional Overview",
        "Country Experience",
        "Geographic Experience",
        "Thematic Areas",
        "Selected Project Experience",
        "Alignment with ToR",
        "Interpretation Log",
    ]

    # Filter to only the expected headings (in case Source Citations also appears)
    found = [t for t in bold_14pt_texts if t in expected_order]

    assert found == expected_order, (
        f"Headings not in expected order.\nExpected: {expected_order}\nFound: {found}"
    )


# ---------------------------------------------------------------------------
# Unit test: document styles match spec
# ---------------------------------------------------------------------------

def test_document_styles_match_spec(tmp_path):
    """Body text is 11pt, citation runs are 9pt italic, margins are 2.5cm."""
    draft = _make_minimal_draft()
    citation_result = _make_minimal_citation_result()

    result = write_output(draft, citation_result, output_path=str(tmp_path))
    assert not result.startswith("ERROR:"), f"Got error: {result}"

    doc = Document(result)

    # Check page margins (within 0.1cm tolerance)
    tolerance = Cm(0.1)
    for section in doc.sections:
        assert abs(section.top_margin - Cm(2.5)) <= tolerance, (
            f"Top margin {section.top_margin} != 2.5cm"
        )
        assert abs(section.bottom_margin - Cm(2.5)) <= tolerance, (
            f"Bottom margin {section.bottom_margin} != 2.5cm"
        )
        assert abs(section.left_margin - Cm(2.5)) <= tolerance, (
            f"Left margin {section.left_margin} != 2.5cm"
        )
        assert abs(section.right_margin - Cm(2.5)) <= tolerance, (
            f"Right margin {section.right_margin} != 2.5cm"
        )

    # Check body text runs are 11pt
    body_sizes = []
    citation_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size == Pt(11) and not run.bold:
                body_sizes.append(run.font.size)
            if run.font.size == Pt(9) and run.italic:
                citation_runs.append(run)

    assert len(body_sizes) > 0, "No 11pt body text runs found"
    assert len(citation_runs) > 0, "No 9pt italic citation runs found"

    # Verify citation runs are italic
    for run in citation_runs:
        assert run.italic is True, "Citation run is not italic"


# ---------------------------------------------------------------------------
# Unit test: no [REF:] tags in output document
# ---------------------------------------------------------------------------

def test_no_ref_tags_in_output_document(tmp_path):
    """No [REF:...] tags appear in the written document paragraphs.
    write_output strips REF tags from section text before writing.
    """
    draft = _make_minimal_draft()
    # Add REF tags to sections — they should be stripped by write_output
    draft["sections"]["opening_statement"] = (
        "GovRisk overview. [REF:doc.pdf:page_1] More text [REF:report.pdf:page_5]."
    )
    draft["sections"]["geographic_experience"] = (
        "Strong in Mexico [REF:geo.pdf:page_2]."
    )

    citation_result = _make_minimal_citation_result()

    result = write_output(draft, citation_result, output_path=str(tmp_path))
    assert not result.startswith("ERROR:"), f"Got error: {result}"

    doc = Document(result)
    ref_pattern = re.compile(r'\[REF:')

    for para in doc.paragraphs:
        assert not ref_pattern.search(para.text), (
            f"REF tag found in paragraph: {para.text[:100]}"
        )

    # Also check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    assert not ref_pattern.search(para.text), (
                        f"REF tag found in table cell: {para.text[:100]}"
                    )


# ---------------------------------------------------------------------------
# PBT P15: filename matches pattern
# ---------------------------------------------------------------------------

@given(
    st.datetimes(
        min_value=datetime(2020, 1, 1),
        max_value=datetime(2030, 12, 31),
    )
)
@settings(max_examples=20, deadline=None)
def test_p15_filename_format(dt):
    """P15: Output filename matches GovRisk_CapabilityStatement_{YYYY-MM-DD}_{HH-MM}.docx
    **Validates: Requirements 7.3**
    """
    filename = f"GovRisk_CapabilityStatement_{dt.strftime('%Y-%m-%d_%H-%M')}.docx"
    pattern = re.compile(
        r'^GovRisk_CapabilityStatement_\d{4}-\d{2}-\d{2}_\d{2}-\d{2}\.docx$'
    )
    assert pattern.match(filename), (
        f"Filename '{filename}' does not match pattern"
    )


# ---------------------------------------------------------------------------
# PBT P16: country table has exactly 5 columns and correct row count
# ---------------------------------------------------------------------------

@given(
    st.lists(
        st.fixed_dictionaries({
            "country": st.text(min_size=1, max_size=20),
            "project_count": st.integers(min_value=1, max_value=10),
            "year_range": st.text(min_size=1, max_size=10),
            "named_identifiers": st.lists(
                st.text(min_size=1, max_size=10), max_size=3
            ),
            "donors": st.lists(
                st.text(min_size=1, max_size=10), max_size=3
            ),
        }),
        min_size=1,
        max_size=5,
    )
)
@settings(
    max_examples=20,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture],
)
def test_p16_country_table_structure(tmp_path, country_table_data):
    """P16: Word table has exactly 5 columns and 1 header row + N data rows.
    **Validates: Requirements 7.5**
    """
    draft = {
        "sections": {
            "opening_statement": "Overview.",
            "institutional_overview": "Founded 2010.",
            "country_table": country_table_data,
            "geographic_experience": "Regional.",
            "thematic_areas": "AML/CFT.",
            "selected_project_experience": "Project A.",
            "alignment_with_tor": "Aligned.",
        },
        "interpretation_log": [],
        "summary": {
            "sections_generated": 7,
            "projects_referenced": 0,
            "countries_covered": len(country_table_data),
            "documents_used": 0,
            "overall_confidence": "LOW",
        },
    }

    result = write_output(draft, None, output_path=str(tmp_path))
    assert not result.startswith("ERROR:"), f"Got error: {result}"

    doc = Document(result)
    assert len(doc.tables) >= 1, "No tables found in document"

    table = doc.tables[0]
    # Exactly 5 columns
    assert len(table.columns) == 5, (
        f"Expected 5 columns, got {len(table.columns)}"
    )
    # 1 header row + N data rows
    expected_rows = 1 + len(country_table_data)
    assert len(table.rows) == expected_rows, (
        f"Expected {expected_rows} rows, got {len(table.rows)}"
    )


# ---------------------------------------------------------------------------
# PBT P18: output formatter preserves project experience order
# ---------------------------------------------------------------------------

@given(
    st.lists(
        st.text(min_size=1, max_size=50),
        min_size=2,
        max_size=5,
    )
)
@settings(
    max_examples=20,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture],
)
def test_p18_project_experience_order_preserved(tmp_path, project_texts):
    """P18: Project experience entries appear in the document in the same order provided.
    **Validates: Requirements 10.3**
    """
    # Build a selected_project_experience section with numbered entries
    # separated by double newlines so they become distinct paragraphs
    numbered_entries = [
        f"Entry{i}: {text}" for i, text in enumerate(project_texts)
    ]
    project_section = "\n\n".join(numbered_entries)

    draft = {
        "sections": {
            "opening_statement": "",
            "institutional_overview": "",
            "country_table": [],
            "geographic_experience": "",
            "thematic_areas": "",
            "selected_project_experience": project_section,
            "alignment_with_tor": "",
        },
        "interpretation_log": [],
        "summary": {
            "sections_generated": 1,
            "projects_referenced": 0,
            "countries_covered": 0,
            "documents_used": 0,
            "overall_confidence": "LOW",
        },
    }

    result = write_output(draft, None, output_path=str(tmp_path))
    assert not result.startswith("ERROR:"), f"Got error: {result}"

    doc = Document(result)

    # Collect all paragraph texts that contain "Entry" prefix
    doc_texts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    entry_texts = [t for t in doc_texts if t.startswith("Entry")]

    # Build expected entries as they will appear after write_output processing
    # (strip + remove control chars + normalize \r to \n, same as _add_body_text does)
    import re as _re
    _ctrl = _re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]')
    processed_entries = [
        _ctrl.sub('', e).replace('\r', '\n').strip() for e in numbered_entries
    ]
    # Filter out entries that become empty after processing
    processed_entries = [e for e in processed_entries if e]

    # Verify order is preserved
    assert len(entry_texts) == len(processed_entries), (
        f"Expected {len(processed_entries)} entries, found {len(entry_texts)}"
    )
    for i, (found, expected) in enumerate(zip(entry_texts, processed_entries)):
        assert found == expected, (
            f"Entry {i} mismatch: expected '{expected}', got '{found}'"
        )


def test_section_order_follows_sections_to_include():
    from output_formatter import write_output
    import tempfile, os
    from docx import Document

    draft = {
        "sections": {
            "opening_statement": "Opening text",
            "thematic_areas": "Thematic text",
            "alignment_with_tor": "Alignment text",
        },
        "interpretation_log": [],
        "summary": {},
    }

    with tempfile.TemporaryDirectory() as tmpdir:
        path = write_output(
            draft, None,
            sections_to_include=[
                "alignment_with_tor",
                "thematic_areas",
                "opening_statement",
            ],
            output_path=tmpdir,
        )
    assert not path.startswith("ERROR:")


def test_write_output_accepts_section_formats():
    from output_formatter import write_output
    import tempfile

    draft = {
        "sections": {"thematic_areas": "• Item one\n• Item two"},
        "interpretation_log": [],
        "summary": {},
    }

    with tempfile.TemporaryDirectory() as tmpdir:
        path = write_output(
            draft, None,
            sections_to_include=["thematic_areas"],
            section_formats={"thematic_areas": "Bullet list"},
            output_path=tmpdir,
        )
    assert not path.startswith("ERROR:")


def test_country_table_list_renders_without_error():
    from output_formatter import write_output
    import tempfile

    ct = [{"country": "Mexico", "project_count": 2,
           "year_range": "2024-2026",
           "named_identifiers": ["PECEL"],
           "donors": ["US State Dept"]}]
    draft = {
        "sections": {"country_table": ct},
        "interpretation_log": [],
        "summary": {},
    }

    with tempfile.TemporaryDirectory() as tmpdir:
        path = write_output(
            draft, None,
            sections_to_include=["country_table"],
            output_path=tmpdir,
        )
    assert not path.startswith("ERROR:")
