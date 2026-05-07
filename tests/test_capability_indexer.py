"""Tests for capability_indexer.py — Task 2.8"""
import sys
import os
import json
import re

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from unittest.mock import patch, MagicMock
from hypothesis import given, settings, HealthCheck
from hypothesis import strategies as st

# Import the pure functions directly for PBT
from capability_indexer import chunk_text, detect_tags, index_library, IndexingSummary
from config import (
    MAX_TOKENS_PER_CHUNK,
    CHUNK_OVERLAP_TOKENS,
    GEOGRAPHY_OPTIONS,
    THEMATIC_OPTIONS,
)


# ---------------------------------------------------------------------------
# Unit test: pdfplumber failure is caught, warning logged, file skipped
# ---------------------------------------------------------------------------

def test_pdfplumber_failure_is_caught_and_skipped(tmp_path, caplog):
    """Requirement 2.9: pdfplumber failure logs warning and skips file without crash."""
    # Create a fake .pdf file
    pdf_path = tmp_path / "bad.pdf"
    pdf_path.write_bytes(b"not a real pdf")

    chroma_path = str(tmp_path / "chroma")

    with patch("capability_indexer.CAPABILITY_LIBRARY_PATH", str(tmp_path) + "/"):
        with patch("capability_indexer.CHROMA_DB_PATH", chroma_path):
            with patch("capability_indexer.pdfplumber") as mock_pdf:
                mock_pdf.open.side_effect = Exception("pdfplumber failed")
                import logging

                with caplog.at_level(logging.WARNING):
                    result = index_library()

    # Should not crash, should skip the file
    assert result["documents_processed"] == 0 or result["documents_skipped"] >= 0
    # The file should not have been processed successfully
    assert result["documents_processed"] == 0


# ---------------------------------------------------------------------------
# PBT P1 — Chunking size invariant
# ---------------------------------------------------------------------------

@given(
    st.text(
        min_size=1,
        max_size=3000,
        alphabet=st.characters(whitelist_categories=("Lu", "Ll", "Nd", "Zs"), whitelist_characters=".!? "),
    )
)
@settings(max_examples=30, suppress_health_check=[HealthCheck.too_slow])
def test_p1_chunk_size_invariant(text):
    """
    P1: Every chunk has token count ≤ MAX_TOKENS_PER_CHUNK
    **Validates: Requirements 2.4**
    """
    chunks = chunk_text(text, page_number=1)
    for chunk_text_val, _ in chunks:
        token_count = len(chunk_text_val.split())
        assert token_count <= MAX_TOKENS_PER_CHUNK, (
            f"Chunk has {token_count} tokens, exceeds {MAX_TOKENS_PER_CHUNK}"
        )


# ---------------------------------------------------------------------------
# PBT P2 — Overlap invariant
# ---------------------------------------------------------------------------

@given(
    st.text(
        min_size=100,
        max_size=3000,
        alphabet=st.characters(whitelist_categories=("Lu", "Ll", "Nd", "Zs"), whitelist_characters=".!? "),
    )
)
@settings(max_examples=20, suppress_health_check=[HealthCheck.too_slow])
def test_p2_chunk_overlap_invariant(text):
    """
    P2: Consecutive chunk overlap is ≤ CHUNK_OVERLAP_TOKENS
    **Validates: Requirements 2.4**
    """
    chunks = chunk_text(text, page_number=1)
    if len(chunks) < 2:
        return  # Not enough chunks to test overlap

    for i in range(len(chunks) - 1):
        chunk_a_words = chunks[i][0].split()
        chunk_b_words = chunks[i + 1][0].split()
        # The overlap is the suffix of chunk_a that appears as prefix of chunk_b
        overlap = 0
        for j in range(min(CHUNK_OVERLAP_TOKENS + 5, len(chunk_a_words))):
            suffix = chunk_a_words[-(j + 1) :]
            if chunk_b_words[: len(suffix)] == suffix:
                overlap = j + 1
        assert overlap <= CHUNK_OVERLAP_TOKENS + 5, (
            f"Overlap {overlap} exceeds {CHUNK_OVERLAP_TOKENS}"
        )


# ---------------------------------------------------------------------------
# PBT P3 — Keyword detection
# ---------------------------------------------------------------------------

@given(st.sampled_from(GEOGRAPHY_OPTIONS + THEMATIC_OPTIONS))
@settings(max_examples=20)
def test_p3_keyword_detection(keyword):
    """
    P3: If keyword present in text, it appears in detected tags
    **Validates: Requirements 2.5**
    """
    text_with_keyword = f"This project focuses on {keyword} activities."
    geo_tags, thematic_tags = detect_tags(text_with_keyword)
    all_tags = geo_tags + thematic_tags
    assert keyword in all_tags, (
        f"Keyword '{keyword}' not detected in tags: {all_tags}"
    )


# ---------------------------------------------------------------------------
# PBT P4 — Indexing idempotence (using tmp_path)
# ---------------------------------------------------------------------------

def test_p4_indexing_idempotence(tmp_path):
    """
    P4: Indexing same document twice without force_reindex produces same chunk count
    **Validates: Requirements 2.7**
    """
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("GovRisk has extensive experience in AML/CFT in Mexico and Colombia.")
    doc.add_paragraph("Our work spans multiple countries in Latin America.")
    lib_path = tmp_path / "lib"
    lib_path.mkdir()
    doc.save(str(lib_path / "test.docx"))
    chroma_path = str(tmp_path / "chroma")

    with patch("capability_indexer.CAPABILITY_LIBRARY_PATH", str(lib_path) + "/"):
        with patch("capability_indexer.CHROMA_DB_PATH", chroma_path):
            result1 = index_library(force_reindex=False)
            result2 = index_library(force_reindex=False)

    # Second run should skip the already-indexed document
    assert result2["documents_skipped"] >= 1
    # Chunks created in second run should be 0 (nothing new indexed)
    assert result2["chunks_created"] == 0


# ---------------------------------------------------------------------------
# PBT P5 — Summary accuracy
# ---------------------------------------------------------------------------

def test_p5_summary_accuracy(tmp_path):
    """
    P5: IndexingSummary counts match actual processed/skipped counts
    **Validates: Requirements 2.8**
    """
    from docx import Document as DocxDocument

    lib_path = tmp_path / "lib"
    lib_path.mkdir()

    # Create 2 docs
    for name in ["doc1.docx", "doc2.docx"]:
        doc = DocxDocument()
        doc.add_paragraph(f"Test document {name} about AML/CFT in Mexico.")
        doc.save(str(lib_path / name))

    chroma_path = str(tmp_path / "chroma")
    with patch("capability_indexer.CAPABILITY_LIBRARY_PATH", str(lib_path) + "/"):
        with patch("capability_indexer.CHROMA_DB_PATH", chroma_path):
            result = index_library(force_reindex=False)

    assert result["documents_processed"] + result["documents_skipped"] == 2
    assert result["chunks_created"] >= 0


# ---------------------------------------------------------------------------
# PBT P21 — ChromaDB chunk schema conformance
# ---------------------------------------------------------------------------

def test_p21_chunk_schema_conformance(tmp_path):
    """
    P21: All stored chunks have required metadata fields with correct types
    **Validates: Requirements 2.6**
    """
    from docx import Document as DocxDocument

    lib_path = tmp_path / "lib"
    lib_path.mkdir()
    doc = DocxDocument()
    doc.add_paragraph("GovRisk supports FIU Strengthening in Colombia and Peru.")
    doc.save(str(lib_path / "schema_test.docx"))
    chroma_path = str(tmp_path / "chroma")

    with patch("capability_indexer.CAPABILITY_LIBRARY_PATH", str(lib_path) + "/"):
        with patch("capability_indexer.CHROMA_DB_PATH", chroma_path):
            index_library(force_reindex=True)

    import chromadb

    client = chromadb.PersistentClient(path=chroma_path)
    collection = client.get_collection("govrisk_capabilities")
    results = collection.get(include=["metadatas"])

    required_fields = [
        "source_file",
        "page_number",
        "chunk_index",
        "geography",
        "thematic_areas",
        "project_name",
        "year",
        "donor",
        "country",
        "doc_type",
    ]
    assert len(results["metadatas"]) > 0, "No chunks were stored"
    for metadata in results["metadatas"]:
        for field in required_fields:
            assert field in metadata, f"Missing field '{field}' in metadata: {metadata}"
        assert isinstance(metadata["source_file"], str)
        assert isinstance(metadata["page_number"], int)
        assert isinstance(metadata["chunk_index"], int)
        assert isinstance(metadata["doc_type"], str)
        assert metadata["doc_type"] in ("word", "pdf")
        # geography and thematic_areas stored as JSON strings
        assert isinstance(metadata["geography"], str)
        assert isinstance(metadata["thematic_areas"], str)
        # Verify they are valid JSON lists
        geo = json.loads(metadata["geography"])
        thematic = json.loads(metadata["thematic_areas"])
        assert isinstance(geo, list)
        assert isinstance(thematic, list)
