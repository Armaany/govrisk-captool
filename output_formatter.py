"""output_formatter.py — Task 7

Writes the capability statement as a formatted .docx file.
"""
import os
import re
import logging
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Pattern to strip inline REF tags from section text before writing
_REF_STRIP_PATTERN = re.compile(r'\[REF:[^\]]+\]')

# Pattern to remove XML-incompatible control characters (except tab, newline, carriage return)
_CONTROL_CHAR_PATTERN = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]')

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Style constants (from SPEC.md §5.7)
# ---------------------------------------------------------------------------

FONT_NAME = "Calibri"
HEADING_SIZE = 14       # section headings: Bold, 14pt
SUBHEADING_SIZE = 12    # subheadings: Bold, 12pt
BODY_SIZE = 11          # body text: 11pt, no bold
CITATION_SIZE = 9       # citations: italic, gray, 9pt
INTERP_LOG_SIZE = 10    # interpretation log: italic, 10pt
CITATION_COLOR = RGBColor(128, 128, 128)
MARGIN = Cm(2.5)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _sanitize(text):
    """Remove XML-incompatible control characters from a string."""
    if not text:
        return text
    return _CONTROL_CHAR_PATTERN.sub('', text)


def _add_heading(doc, text):
    """Add a bold 14pt Calibri section heading paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(HEADING_SIZE)
    run.font.name = FONT_NAME
    return para


def _add_body_text(doc, text):
    """Add body text paragraphs (split on \\n\\n), 11pt Calibri, no bold.
    REF tags and XML-incompatible control characters are stripped before writing.
    """
    if not text or not text.strip():
        return
    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        # Strip inline REF tags from display text
        para_text = _REF_STRIP_PATTERN.sub('', para_text).strip()
        # Remove XML-incompatible control characters
        para_text = _sanitize(para_text)
        # Normalize carriage returns to newlines
        para_text = para_text.replace('\r', '\n')
        if para_text:
            para = doc.add_paragraph()
            run = para.add_run(para_text)
            run.bold = False
            run.font.size = Pt(BODY_SIZE)
            run.font.name = FONT_NAME


def _add_citation_line(doc, citation_text):
    """Add an italic gray 9pt Calibri citation line."""
    para = doc.add_paragraph()
    run = para.add_run(citation_text)
    run.italic = True
    run.font.size = Pt(CITATION_SIZE)
    run.font.name = FONT_NAME
    run.font.color.rgb = CITATION_COLOR
    return para


def _add_country_table(doc, country_table):
    """Render country_table as a Word table with 5 columns."""
    if not country_table:
        return

    headers = ["Country", "Projects", "Years", "Named Identifiers", "Donors"]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = FONT_NAME

    # Data rows
    for entry in country_table:
        row_cells = table.add_row().cells
        row_cells[0].text = _sanitize(str(entry.get("country", "")))
        row_cells[1].text = _sanitize(str(entry.get("project_count", "")))
        row_cells[2].text = _sanitize(str(entry.get("year_range", "")))
        row_cells[3].text = _sanitize(", ".join(entry.get("named_identifiers", [])))
        row_cells[4].text = _sanitize(", ".join(entry.get("donors", [])))
        # Set font for all data cells
        for cell in row_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(BODY_SIZE)


def _add_horizontal_rule(doc):
    """Add a horizontal rule paragraph."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '808080')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_interpretation_log(doc, interpretation_log):
    """Add the interpretation log section with a horizontal rule separator."""
    _add_horizontal_rule(doc)

    heading_para = doc.add_paragraph()
    run = heading_para.add_run("Interpretation Log")
    run.bold = True
    run.font.size = Pt(HEADING_SIZE)
    run.font.name = FONT_NAME

    for entry in interpretation_log:
        text = (
            f"Section: {entry.get('section', '')}\n"
            f"Inference: {entry.get('inference_made', '')}\n"
            f"Source: {entry.get('source_used', '')}\n"
            f"Gap: {entry.get('gap_flagged') or 'None'}\n"
            f"Confidence: {entry.get('confidence', 'LOW')}"
        )
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.italic = True
        run.font.size = Pt(INTERP_LOG_SIZE)
        run.font.name = FONT_NAME


def _add_citations_section(doc, citation_result):
    """Add the Source Citations section from citation_result."""
    paragraphs = citation_result.get("paragraphs", []) if citation_result else []
    if not paragraphs:
        return

    heading_para = doc.add_paragraph()
    run = heading_para.add_run("Source Citations")
    run.bold = True
    run.font.size = Pt(HEADING_SIZE)
    run.font.name = FONT_NAME

    for para_data in paragraphs:
        citations = para_data.get("citations", [])
        if citations:
            from citation_tagger import format_citation
            for cit in citations:
                citation_text = format_citation(cit["filename"], cit["page"])
                _add_citation_line(doc, citation_text)


def _add_bullet_list(doc, text):
    """Render text as a bullet list in the document."""
    if not text or not text.strip():
        return
    lines = text.split('\n')
    for line in lines:
        line = _REF_STRIP_PATTERN.sub('', line).strip()
        line = _sanitize(line)
        if not line:
            continue
        if line.startswith(('•', '-', '*')):
            line = line.lstrip('•-* ').strip()
            try:
                para = doc.add_paragraph(style='List Bullet')
            except Exception:
                para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = Pt(BODY_SIZE)
        else:
            _add_body_text(doc, line)


def _add_project_cards(doc, text):
    """Render text as project card blocks separated by horizontal rules."""
    if not text or not text.strip():
        return
    blocks = text.split('\n\n')
    for block in blocks:
        block = _REF_STRIP_PATTERN.sub('', block).strip()
        block = _sanitize(block)
        if not block:
            continue
        _add_horizontal_rule(doc)
        _add_body_text(doc, block)


# ---------------------------------------------------------------------------
# Public interface
# ---------------------------------------------------------------------------

def write_output(
    generated_draft,
    citation_result,
    output_language="English",
    sections_to_include=None,
    section_formats=None,
    output_path=None,
):
    """
    Write the capability statement as a formatted .docx file.

    Args:
        generated_draft:     GeneratedDraft dict from draft_generator.
        citation_result:     citation_result dict from citation_tagger.
        output_language:     Language string (informational, not used for formatting).
        sections_to_include: Optional list of section keys to include; None = all.
        output_path:         Directory to write the file; None = use OUTPUT_PATH from config.

    Returns:
        The file path string on success.
        An error message string starting with "ERROR:" on failure.
        Never raises exceptions.
    """
    try:
        from config import OUTPUT_PATH

        section_formats = section_formats or {}

        if output_path is None:
            output_path = OUTPUT_PATH

        os.makedirs(output_path, exist_ok=True)

        sections = generated_draft.get("sections", {}) if generated_draft else {}
        interpretation_log = (
            generated_draft.get("interpretation_log", []) if generated_draft else []
        )

        doc = Document()

        # Set page margins
        for sec in doc.sections:
            sec.top_margin = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin = Cm(2.5)
            sec.right_margin = Cm(2.5)

        # Section order — respects sections_to_include order
        _DISPLAY_NAMES = {
            "opening_statement":            "Opening Statement",
            "institutional_overview":       "Institutional Overview",
            "country_table":                "Country Experience",
            "geographic_experience":        "Geographic Experience",
            "thematic_areas":               "Thematic Areas",
            "selected_project_experience":  "Selected Project Experience",
            "alignment_with_tor":           "Alignment with ToR",
        }

        _DEFAULT_ORDER = [
            "opening_statement",
            "institutional_overview",
            "country_table",
            "geographic_experience",
            "thematic_areas",
            "selected_project_experience",
            "alignment_with_tor",
        ]

        ordered_keys = sections_to_include if sections_to_include else _DEFAULT_ORDER
        section_map = [
            (_DISPLAY_NAMES.get(k, k.replace("_", " ").title()), k)
            for k in ordered_keys
        ]

        for heading_text, section_key in section_map:
            format_choice = section_formats.get(section_key, "Narrative")

            _add_heading(doc, heading_text)

            if section_key == "country_table":
                ct = sections.get("country_table")
                if isinstance(ct, list) and ct:
                    _add_country_table(doc, ct)
                elif isinstance(ct, str) and ct.strip():
                    _add_body_text(doc, ct)
            elif format_choice == "Bullet list":
                _add_bullet_list(doc, sections.get(section_key, ""))
            elif format_choice == "Project cards":
                _add_project_cards(doc, sections.get(section_key, ""))
            else:
                _add_body_text(doc, sections.get(section_key, ""))

        # Interpretation log
        _add_interpretation_log(doc, interpretation_log)

        # Citations
        _add_citations_section(doc, citation_result)

        # Generate filename and save
        filename = (
            f"GovRisk_CapabilityStatement_"
            f"{datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx"
        )
        filepath = os.path.join(output_path, filename)
        doc.save(filepath)

        return filepath

    except Exception as exc:
        return f"ERROR: Failed to write output document: {exc}"
